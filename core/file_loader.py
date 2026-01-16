"""
Document loading for various file formats.
"""

import logging
from typing import Optional, Dict, Any
from pathlib import Path
import mimetypes

logger = logging.getLogger(__name__)


class FileLoader:
    """
    Loads and extracts text from various file formats.
    
    Supported formats:
    - Text: .txt, .md
    - PDF: .pdf (with optional OCR)
    - Word: .docx
    - Images: .png, .jpg, .jpeg (with OCR)
    """
    
    def __init__(self, enable_ocr: bool = True):
        """
        Initialize the file loader.
        
        Args:
            enable_ocr: Whether to enable OCR for PDFs and images
        """
        self.enable_ocr = enable_ocr
        self._check_dependencies()
        
        logger.info(f"FileLoader initialized (OCR: {enable_ocr})")
    
    def _check_dependencies(self):
        """Check if optional dependencies are available."""
        self.has_pdf = False
        self.has_docx = False
        self.has_ocr = False
        
        try:
            import pypdf
            self.has_pdf = True
        except ImportError:
            logger.warning("pypdf not installed - PDF support disabled")
        
        try:
            import docx
            self.has_docx = True
        except ImportError:
            logger.warning("python-docx not installed - DOCX support disabled")
        
        if self.enable_ocr:
            try:
                import pytesseract
                from PIL import Image
                self.has_ocr = True
            except ImportError:
                logger.warning("pytesseract/PIL not installed - OCR disabled")
    
    def load(self, file_path: str) -> Optional[Dict[str, Any]]:
        """
        Load a file and extract its text content.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary with 'text', 'metadata', or None if failed
        """
        path = Path(file_path)
        
        if not path.exists():
            logger.error(f"File not found: {file_path}")
            return None
        
        if not path.is_file():
            logger.warning(f"Not a file: {file_path}")
            return None
        
        # Get file extension
        ext = path.suffix.lower()
        
        # Dispatch to appropriate loader
        loaders = {
            '.txt': self._load_text,
            '.md': self._load_text,
            '.pdf': self._load_pdf,
            '.docx': self._load_docx,
            '.png': self._load_image,
            '.jpg': self._load_image,
            '.jpeg': self._load_image,
        }
        
        loader = loaders.get(ext)
        if loader is None:
            logger.warning(f"Unsupported file type: {ext}")
            return None
        
        try:
            text = loader(path)
            
            if not text or not text.strip():
                logger.warning(f"No text extracted from: {file_path}")
                return None
            
            # Build metadata
            metadata = {
                'file_path': str(path.absolute()),
                'file_name': path.name,
                'file_type': ext[1:],  # Remove leading dot
                'file_size': path.stat().st_size,
                'modified_time': path.stat().st_mtime,
            }
            
            logger.debug(f"Loaded {len(text)} chars from {path.name}")
            
            return {
                'text': text,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Error loading {file_path}: {e}")
            return None
    
    def _load_text(self, path: Path) -> str:
        """Load plain text or markdown files."""
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        # Fallback: read with error handling
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    def _load_pdf(self, path: Path) -> str:
        """Load PDF files with optional OCR."""
        if not self.has_pdf:
            raise ImportError("pypdf not installed")
        
        from pypdf import PdfReader
        
        text_parts = []
        
        try:
            reader = PdfReader(str(path))
            
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                
                # If page has little text and OCR is enabled, try OCR
                if self.enable_ocr and self.has_ocr and len(page_text.strip()) < 50:
                    logger.debug(f"Attempting OCR on page {page_num + 1}")
                    # Note: Full OCR implementation would require converting PDF page to image
                    # This is a placeholder - production code should handle this
                
                text_parts.append(page_text)
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            raise
    
    def _load_docx(self, path: Path) -> str:
        """Load Word documents."""
        if not self.has_docx:
            raise ImportError("python-docx not installed")
        
        from docx import Document
        
        try:
            doc = Document(str(path))
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    text_parts.append(row_text)
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
            raise
    
    def _load_image(self, path: Path) -> str:
        """Load images with OCR."""
        if not self.enable_ocr or not self.has_ocr:
            logger.warning("OCR not available for image")
            return ""
        
        try:
            import pytesseract
            from PIL import Image
            
            image = Image.open(str(path))
            text = pytesseract.image_to_string(image)
            
            return text
            
        except Exception as e:
            logger.error(f"Error performing OCR: {e}")
            raise
    
    @staticmethod
    def is_supported(file_path: str) -> bool:
        """
        Check if a file type is supported.
        
        Args:
            file_path: Path to the file
            
        Returns:
            True if file type is supported
        """
        ext = Path(file_path).suffix.lower()
        supported = {'.txt', '.md', '.pdf', '.docx', '.png', '.jpg', '.jpeg'}
        return ext in supported
    
    @staticmethod
    def get_supported_extensions() -> list:
        """Get list of supported file extensions."""
        return ['.txt', '.md', '.pdf', '.docx', '.png', '.jpg', '.jpeg']